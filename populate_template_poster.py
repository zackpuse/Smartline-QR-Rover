import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
import os

def populate_poster():
    template_path = "Template Poster A4 eProsiding MIPAC TVET 2026.docx"
    output_path = "Poster_A4_eProsiding_SmartLine_QR_Rover_MIPACTVET2026.docx"
    
    # 1. Copy template file
    shutil.copy(template_path, output_path)
    
    doc = docx.Document(output_path)
    table = doc.tables[0]

    # Helper styling function
    def add_sec_heading(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF) # Navy Blue
        return p

    def add_p(cell, text, bold_prefix="", space_after=4, size=9.5):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.font.bold = True
            r1.font.size = Pt(size)
            r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        r2 = p.add_run(text)
        r2.font.size = Pt(size)
        r2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    def add_bullet(cell, bold_prefix, text, size=9):
        p = cell.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(size)
        r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        r2 = p.add_run(text)
        r2.font.size = Pt(size)
        r2.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # ----------------------------------------------------
    # ROW 0: HEADER (TITLE & META)
    # ----------------------------------------------------
    cell_r0c0 = table.rows[0].cells[0]
    cell_r0c1 = table.rows[0].cells[1]

    # Clear placeholder text in Row 0
    cell_r0c0.text = ""
    cell_r0c1.text = ""

    # Populate Header in cell_r0c0
    p_title = cell_r0c0.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run("TAJUK INOVASI MODUL PdP INTERAKTIF: SIMULATOR DIAGNOSTIK ADAS AUTOMOTIF MENGGUNAKAN SMARTLINE QR ROVER")
    r_t.font.bold = True
    r_t.font.size = Pt(14)
    r_t.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p_title.paragraph_format.space_after = Pt(4)

    p_meta = cell_r0c0.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_m = p_meta.add_run("Kategori: AI DigiTeach | Nama Kumpulan: Kumpulan Innovate TVET ADTEC KL\nInstitusi: ADTEC JTM Kampus Kuala Lumpur\nNama Ketua Projek: Mohd Fadli bin Mohd Tahir | Ahli: Noorazura Binti Mat Shadan, Norakma Binti Nawawi\nE-mel: mohdfadli@jtm.gov.my")
    r_m.font.size = Pt(9.5)
    r_m.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    p_meta.paragraph_format.space_after = Pt(4)

    # Merge Row 0 cells if not already merged
    try:
        cell_r0c0.merge(cell_r0c1)
    except Exception as e:
        print("Row 0 merge status:", e)

    # ----------------------------------------------------
    # ROW 1, CELL 0: 1. LATAR BELAKANG & 2. OBJEKTIF
    # ----------------------------------------------------
    cell_r1c0 = table.rows[1].cells[0]
    cell_r1c0.text = "" # Clear placeholder

    add_sec_heading(cell_r1c0, "1. LATAR BELAKANG / PENYATAAN MASALAH")
    add_p(cell_r1c0, "Projek ini dibangunkan bagi mengatasi masalah kos tinggi penyediaan kenderaan sebenar ber-ADAS serta keselamatan pengendalian sistem bervoltan tinggi yang dihadapi oleh pelatih TVET Automotif. Kaedah sedia ada didapati bergantung sepenuhnya kepada manual teknikal dan video tanpa latihan amali langsung, sekali gus menjejaskan kefahaman praktikal logik kawalan dan pengaturcaraan sensor ADAS.")

    add_sec_heading(cell_r1c0, "2. OBJEKTIF INOVASI")
    add_p(cell_r1c0, "Projek ini bertujuan untuk:")
    add_bullet(cell_r1c0, "i. Simulator Robotik ADAS: ", "Membangunkan platform SmartLine QR Rover menggunakan AI Vision (HuskyLens) dan sensor ultrasonik bagi mensimulasikan ADAS.")
    add_bullet(cell_r1c0, "ii. Modul PdP Digital NOSS: ", "Menyediakan modul PdP digital interaktif yang dipetakan secara langsung kepada standard NOSS G452-011-4:2025 (ADAS Diagnosis).")
    add_bullet(cell_r1c0, "iii. Latihan Amali Kos Rendah: ", "Menyediakan pendedahan amali berskala mikro dan kos rendah sebelum pelatih mengendalikan kenderaan sebenar.")

    # ----------------------------------------------------
    # ROW 1, CELL 1: 3. PENERANGAN INOVASI & GAMBAR PROTOTAIP
    # ----------------------------------------------------
    cell_r1c1 = table.rows[1].cells[1]
    cell_r1c1.text = "" # Clear placeholder

    add_sec_heading(cell_r1c1, "3. PENERANGAN INOVASI")
    add_p(cell_r1c1, "SmartLine QR Rover ialah platform robotik pembelajaran interaktif yang dibangunkan khusus untuk mensimulasikan logik kamera pengesanan visual dan sensor perlanggaran kenderaan autonomi ADAS dalam persekitaran mikro yang selamat dan mudah dikawal.")
    
    add_p(cell_r1c1, "Fungsi Utama:", bold_prefix="", space_after=2)
    add_bullet(cell_r1c1, "i. Pengecaman Kod QR (AI Vision): ", "Mensimulasikan modul kamera Traffic Sign Recognition (TSR).")
    add_bullet(cell_r1c1, "ii. Pengesanan Halangan Ultrasonik: ", "Mensimulasikan sistem Automatic Emergency Braking (AEB).")
    add_bullet(cell_r1c1, "iii. Navigasi Autonomi (Line Following): ", "Membimbing rover mengikut trek garisan secara automatik.")
    add_bullet(cell_r1c1, "iv. Kawalan Henti/Resume Automatik: ", "Berhenti apabila mengesan objek dan menyambung pergerakan selepas laluan bersih.")

    # Insert Image 1 (SmartLine QR Rover Prototype)
    img1_path = r'C:\Users\palie\.gemini\antigravity-ide\brain\e4e3e780-2bda-4fb6-9a09-94ba09225e31\smartline_qr_rover.png'
    if os.path.exists(img1_path):
        p_img1 = cell_r1c1.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.paragraph_format.space_before = Pt(6)
        p_img1.paragraph_format.space_after = Pt(2)
        r_img1 = p_img1.add_run()
        r_img1.add_picture(img1_path, width=Inches(3.4))

        p_cap1 = cell_r1c1.add_paragraph()
        p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap1 = p_cap1.add_run("Rajah 1. Prototaip dan Aliran Kerja Operasi SmartLine QR Rover")
        r_cap1.font.italic = True
        r_cap1.font.bold = True
        r_cap1.font.size = Pt(8.5)
        r_cap1.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # ----------------------------------------------------
    # ROW 2, CELL 0: 4. METODOLOGI / PROSES PEMBANGUNAN
    # ----------------------------------------------------
    cell_r2c0 = table.rows[2].cells[0]
    cell_r2c0.text = "" # Clear placeholder

    add_sec_heading(cell_r2c0, "4. METODOLOGI / PROSES PEMBANGUNAN")
    add_bullet(cell_r2c0, "Fasa 1: Mengenal Pasti Masalah: ", "Menilai jurang latihan amali ADAS akibat kos kenderaan tinggi dan risiko keselamatan.")
    add_bullet(cell_r2c0, "Fasa 2: Reka Bentuk: ", "Merangka casis akrilik pacuan 4-roda dan menyepadukan sensor AI Vision, ultrasonik, serta IR garisan.")
    add_bullet(cell_r2c0, "Fasa 3: Pembangunan: ", "Mengatur cara mikropengawal Arduino (ECU) dan membangunkan modul PdP digital Google Site selari NOSS G452-011-4:2025.")
    add_bullet(cell_r2c0, "Fasa 4: Pengujian: ", "Ujian rintis kerosakan sensor terpandu dan tindak balas arahan kod QR bersama pelatih.")
    add_bullet(cell_r2c0, "Fasa 5: Penambahbaikan: ", "Memurnikan soalan refleksi dan mengintegrasikan perancangan sistem pelaporan Diagnostic Trouble Code (DTC).")

    # ----------------------------------------------------
    # ROW 2, CELL 1: 5. NILAI KELEBIHAN & 6. ETIKA AI
    # ----------------------------------------------------
    cell_r2c1 = table.rows[2].cells[1]
    cell_r2c1.text = "" # Clear placeholder

    add_sec_heading(cell_r2c1, "5. NILAI DAN KELEBIHAN INOVASI")
    
    # Table inside Cell 2,1
    t_nilai = cell_r2c1.add_table(rows=7, cols=2)
    t_nilai.autofit = False
    
    headers = ["Komponen", "Penerangan Ringkas"]
    for i, h in enumerate(headers):
        c = t_nilai.rows[0].cells[i]
        c.text = h
        p = c.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(8.5)
        p.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    data_nilai = [
        ("Kebaharuan", "Menterjemahkan konsep diagnosis ADAS kepada persekitaran mikro yang selamat menggunakan AI Vision plug-and-play."),
        ("Kebergunaan", "Membantu pelatih menguasai logik diagnostik C05-W02 & C05-W03 secara hands-on tanpa risiko bervoltan tinggi."),
        ("Kebolehlaksanaan", "Litar pendawaian, kod sumber, dan spesifikasi komponen dipaparkan di Google Site untuk dibina semula."),
        ("Keberkesanan Kos", "Kos pembinaan kit jauh lebih rendah berbanding pemerolehan sistem ADAS atau kenderaan baharu."),
        ("Kelestarian", "Penggunaan bateri boleh cas semula serta komponen modular yang mudah diselenggara dan disokong digital."),
        ("Kebolehskalaan", "Berpotensi diperluaskan kepada simulasi Radar & LiDAR (C05-W01, C05-W04) serta pengkomersialan kit STEM/TVET.")
    ]

    for idx, (k, v) in enumerate(data_nilai, start=1):
        row = t_nilai.rows[idx]
        row.cells[0].text = k
        row.cells[1].text = v
        for c_i, cell_obj in enumerate(row.cells):
            p = cell_obj.paragraphs[0]
            p.runs[0].font.size = Pt(8)
            if c_i == 0:
                p.runs[0].font.bold = True

    add_sec_heading(cell_r2c1, "6. PENGGUNAAN AI DAN ETIKA AI")
    add_p(cell_r2c1, "Claude (Anthropic), ChatGPT, dan Google Flow (Veo 3.1).", bold_prefix="Alat AI: ", size=8.5)
    add_p(cell_r2c1, "AI digunakan untuk pemetaan kurikulum NOSS, penyusunan skrip PdP, dan penjanaan segmen visual konsep. Semua kandungan teknikal, logik pengaturcaraan, dan pengujian fizikal prototaip disemak serta disahkan secara manual oleh penulis.", bold_prefix="Peranan & Pengesahan: ", size=8.5)

    # ----------------------------------------------------
    # ROW 3, CELL 0: 7. HASIL PENGUJIAN & 8. IMPAK INOVASI
    # ----------------------------------------------------
    cell_r3c0 = table.rows[3].cells[0]
    cell_r3c0.text = "" # Clear placeholder

    add_sec_heading(cell_r3c0, "7. HASIL PENGUJIAN / VALIDASI / BUKTI")
    add_p(cell_r3c0, "Hasil pengujian rintis menunjukkan peningkatan kefahaman logik sistem ADAS yang ketara. Pelatih yang pada mulanya sukar membayangkan logik tindak balas AEB menerusi manual teknikal berjaya membuat deduksi logik diagnostik kerosakan pada mikropengawal secara automatik semasa simulasi kegagalan sensor.")

    # Insert Image 2 (Chart)
    img2_path = os.path.join('extracted_media', 'graf_hasil_pengujian.png')
    if os.path.exists(img2_path):
        p_img2 = cell_r3c0.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.paragraph_format.space_before = Pt(4)
        p_img2.paragraph_format.space_after = Pt(2)
        r_img2 = p_img2.add_run()
        r_img2.add_picture(img2_path, width=Inches(3.3))

        p_cap2 = cell_r3c0.add_paragraph()
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap2 = p_cap2.add_run("Rajah 2. Carta Hasil Pengujian Rintis Perbandingan Kefahaman Pelatih")
        r_cap2.font.italic = True
        r_cap2.font.bold = True
        r_cap2.font.size = Pt(8)
        r_cap2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    add_sec_heading(cell_r3c0, "8. IMPAK INOVASI")
    t_impak = cell_r3c0.add_table(rows=6, cols=2)
    t_impak.autofit = False
    
    for i, h in enumerate(["Bidang Impak", "Penerangan"]):
        c = t_impak.rows[0].cells[i]
        c.text = h
        p = c.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(8.5)
        p.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    data_impak = [
        ("Akademik / Pembelajaran", "Merapatkan jurang teori dan praktikal logik kawalan ADAS sebelum latihan kenderaan sebenar."),
        ("Operasi / Institusi", "Membantu pensyarah mencapai standard NOSS G452-011-4:2025 Level 4 secara kos efektif."),
        ("Komuniti / Industri", "Menjadi kit latihan asas berpotensi tinggi untuk institusi kemahiran (Politeknik, ILP, ADTEC) dan vokasional."),
        ("Pengkomersialan", "Berpotensi dikomersialkan sebagai Kit Pendidikan STEM/TVET interaktif berserta perkakasan dan perisian."),
        ("Ekosistem TVET", "Menyokong agenda Industri 4.0 dan transformasi TVET negara dengan pendedahan sensor AI autonomi.")
    ]

    for idx, (k, v) in enumerate(data_impak, start=1):
        row = t_impak.rows[idx]
        row.cells[0].text = k
        row.cells[1].text = v
        for c_i, cell_obj in enumerate(row.cells):
            p = cell_obj.paragraphs[0]
            p.runs[0].font.size = Pt(8)
            if c_i == 0:
                p.runs[0].font.bold = True

    # ----------------------------------------------------
    # ROW 3, CELL 1: 9. KESIMPULAN & 10. RUJUKAN & KOD QR
    # ----------------------------------------------------
    cell_r3c1 = table.rows[3].cells[1]
    cell_r3c1.text = "" # Clear placeholder

    add_sec_heading(cell_r3c1, "9. KESIMPULAN")
    add_p(cell_r3c1, "SmartLine QR Rover berjaya membantu menyelesaikan masalah kekangan kos dan risiko keselamatan latihan diagnostik ADAS. Inovasi ini memberi manfaat tinggi dari aspek kefahaman logik kawalan dan berpotensi diperluaskan secara menyeluruh dalam ekosistem TVET Automotif.")

    add_sec_heading(cell_r3c1, "10. RUJUKAN DAN BAHAN SOKONGAN")
    add_p(cell_r3c1, "Jabatan Pembangunan Kemahiran. (2025). National Occupational Skills Standard (NOSS) G452-011-4:2025 Automotive Electrical Diagnostic Level 4. Kementerian Sumber Manusia Malaysia.", size=8.5)

    add_p(cell_r3c1, "Imbas Kod QR di bawah untuk mengakses modul PdP interaktif, manual penggunaan, dan bukti pelaksanaan:", bold_prefix="Modul PdP Digital: ", size=8.5)

    # Insert Image 3 (QR Code)
    img3_path = os.path.join('extracted_media', 'qr_pdp_module.png')
    if os.path.exists(img3_path):
        p_img3 = cell_r3c1.add_paragraph()
        p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img3.paragraph_format.space_before = Pt(4)
        p_img3.paragraph_format.space_after = Pt(2)
        r_img3 = p_img3.add_run()
        r_img3.add_picture(img3_path, width=Inches(1.5))

        p_cap3 = cell_r3c1.add_paragraph()
        p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap3 = p_cap3.add_run("Kod QR Modul PdP Digital SmartLine QR Rover\n")
        r_cap3.font.bold = True
        r_cap3.font.size = Pt(8.5)

        r_url = p_cap3.add_run("Pautan: https://sites.google.com/jtm.gov.my/smartlineqrrover/home")
        r_url.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        r_url.font.underline = True
        r_url.font.size = Pt(8.5)

    doc.save(output_path)
    print(f"Successfully populated poster template and created: {output_path}")

if __name__ == "__main__":
    populate_poster()
