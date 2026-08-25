import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def build_docx():
    doc = docx.Document()
    
    # Page Setup (A4, 2.0 cm margins)
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Styles setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Helper functions
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        p.paragraph_format.space_after = Pt(4)

    def add_meta(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.font.bold = True
            r1.font.size = Pt(10.5)
            r1.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(2)

    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)

    def add_p(text, bold_prefix="", space_after=6):
        p = doc.add_paragraph()
        if bold_prefix:
            r1 = p.add_run(bold_prefix)
            r1.font.bold = True
            r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15

    def add_bullet(bold_txt, text):
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(bold_txt)
        r1.font.bold = True
        r2 = p.add_run(text)
        p.paragraph_format.space_after = Pt(3)

    # --- TITLE & HEADER ---
    add_title("TAJUK INOVASI MODUL PdP INTERAKTIF: SIMULATOR DIAGNOSTIK ADAS AUTOMOTIF MENGGUNAKAN SMARTLINE QR ROVER")
    add_meta("AI DigiTeach | Kumpulan Innovate TVET ADTEC KL | ADTEC JTM Kampus Kuala Lumpur")
    add_meta("Mohd Fadli bin Mohd Tahir (Ketua Project), Noorazura Binti Mat Shadan, Norakma Binti Nawawi")
    add_meta("mohdfadli@jtm.gov.my", "E-mel Ketua Projek: ")
    add_meta("Logo ADTEC JTM Kampus Kuala Lumpur & JTM", "Logo: ")

    p_divider = doc.add_paragraph()
    p_divider.paragraph_format.space_after = Pt(12)

    # --- 1. LATAR BELAKANG ---
    add_heading("1. LATAR BELAKANG / PENYATAAN MASALAH")
    add_p("Projek ini dibangunkan bagi mengatasi masalah kos tinggi penyediaan kenderaan sebenar ber-ADAS serta keselamatan pengendalian sistem bervoltan tinggi yang dihadapi oleh pelatih TVET Automotif. Kaedah sedia ada didapati bergantung sepenuhnya kepada manual teknikal dan video tanpa latihan amali langsung, sekali gus menjejaskan kefahaman praktikal logik kawalan dan pengaturcaraan sensor ADAS.")

    # --- 2. OBJEKTIF ---
    add_heading("2. OBJEKTIF INOVASI")
    add_p("Projek ini bertujuan untuk:")
    add_bullet("i. Membangunkan platform simulator robotik berautonomi SmartLine QR Rover ", "menggunakan AI Vision (HuskyLens) dan sensor ultrasonik bagi mensimulasikan ADAS.")
    add_bullet("ii. Menyediakan modul PdP digital interaktif ", "yang dipetakan secara langsung kepada standard NOSS G452-011-4:2025 (ADAS Diagnosis).")
    add_bullet("iii. Menyediakan pendedahan amali berskala mikro ", "dan kos rendah sebelum pelatih mengendalikan kenderaan sebenar.")

    # --- 3. PENERANGAN INOVASI ---
    add_heading("3. PENERANGAN INOVASI")
    add_p("SmartLine QR Rover ialah platform robotik pembelajaran interaktif yang dibangunkan khusus untuk mensimulasikan logik kamera pengesanan visual dan sensor perlanggaran kenderaan autonomi ADAS dalam persekitaran mikro yang selamat dan mudah dikawal.")
    
    add_p("Fungsi Utama Inovasi:")
    add_bullet("i. Pengecaman Kod QR (AI Vision): ", "Mensimulasikan modul kamera Traffic Sign Recognition (TSR).")
    add_bullet("ii. Pengesanan Halangan Ultrasonik: ", "Mensimulasikan sistem pembrekan kecemasan automatik (Automatic Emergency Braking / AEB).")
    add_bullet("iii. Navigasi Autonomi (Line Following): ", "Membimbing rover mengikut trek garisan secara automatik.")
    add_bullet("iv. Kawalan Henti/Resum Automatik: ", "Berhenti apabila mengesan objek dan menyambung pergerakan secara automatik selepas laluan bersih.")

    # IMAGE 1: PROTOTYPE
    img1_path = r'C:\Users\palie\.gemini\antigravity-ide\brain\e4e3e780-2bda-4fb6-9a09-94ba09225e31\smartline_qr_rover.png'
    if os.path.exists(img1_path):
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img1 = p_img1.add_run()
        run_img1.add_picture(img1_path, width=Inches(5.8))
        
        p_cap1 = doc.add_paragraph()
        p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap1 = p_cap1.add_run("Rajah 1. Prototaip dan Aliran Kerja Operasi SmartLine QR Rover")
        r_cap1.font.italic = True
        r_cap1.font.bold = True
        r_cap1.font.size = Pt(9.5)
        r_cap1.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # --- 4. METODOLOGI ---
    add_heading("4. METODOLOGI / PROSES PEMBANGUNAN")
    add_bullet("Fasa 1: Mengenal Pasti Masalah: ", "Menilai jurang latihan amali ADAS akibat kos kenderaan tinggi dan risiko keselamatan.")
    add_bullet("Fasa 2: Reka Bentuk: ", "Merangka casis akrilik pacuan 4-roda dan menyepadukan sensor AI Vision, ultrasonik, serta IR garisan.")
    add_bullet("Fasa 3: Pembangunan: ", "Mengatur cara mikropengawal Arduino (ECU) dan membangunkan modul PdP digital Google Site selari NOSS G452-011-4:2025.")
    add_bullet("Fasa 4: Pengujian: ", "Ujian rintis kerosakan sensor terpandu dan tindak balas arahan kod QR bersama pelatih.")
    add_bullet("Fasa 5: Penambahbaikan: ", "Memurnikan soalan refleksi dan mengintegrasikan perancangan sistem pelaporan Diagnostic Trouble Code (DTC).")

    # --- 5. NILAI DAN KELEBIHAN ---
    add_heading("5. NILAI DAN KELEBIHAN INOVASI")
    
    table_nilai = doc.add_table(rows=7, cols=2)
    table_nilai.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_nilai.autofit = False

    headers_n = ["Komponen", "Penerangan"]
    hdr_cells_n = table_nilai.rows[0].cells
    for i, title in enumerate(headers_n):
        hdr_cells_n[i].text = title
        set_cell_background(hdr_cells_n[i], "1E40AF")
        set_cell_margins(hdr_cells_n[i], 120, 120, 150, 150)
        p = hdr_cells_n[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10.5)

    data_nilai = [
        ("Kebaharuan (Novelty)", "Menterjemahkan konsep diagnosis ADAS (skala penuh) kepada persekitaran mikro yang selamat menggunakan AI Vision plug-and-play."),
        ("Kebergunaan (Usefulness)", "Membantu pelatih menguasai logik diagnostik C05-W02 & C05-W03 secara hands-on tanpa risiko keselamatan bervoltan tinggi."),
        ("Kebolehlaksanaan (Replicability)", "Litar pendawaian, kod sumber, dan spesifikasi komponen dipaparkan di Google Site untuk dibina semula oleh institusi TVET lain."),
        ("Keberkesanan Kos (Cost Effectiveness)", "Kos pembinaan kit jauh lebih rendah berbanding pemerolehan sistem ADAS atau kenderaan baharu untuk latihan awal."),
        ("Kelestarian (Sustainability)", "Penggunaan bateri boleh cas semula serta komponen modular yang mudah diselenggara dan disokong platform digital."),
        ("Kebolehskalaan (Scalability)", "Berpotensi diperluaskan kepada simulasi Radar & LiDAR (C05-W01, C05-W04) serta pengkomersialan kit STEM/TVET.")
    ]

    col_widths_n = [Inches(2.2), Inches(4.5)]
    for row_idx, data in enumerate(data_nilai, start=1):
        row_cells = table_nilai.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for c_idx in range(2):
            row_cells[c_idx].text = data[c_idx]
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], 100, 100, 150, 150)
            p = row_cells[c_idx].paragraphs[0]
            p.runs[0].font.size = Pt(10)
            if c_idx == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    for row in table_nilai.rows:
        for i, w in enumerate(col_widths_n):
            row.cells[i].width = w

    # --- 6. PENGGUNAAN AI DAN ETIKA AI ---
    add_heading("6. PENGGUNAAN AI DAN ETIKA AI")
    add_p("Alat AI: ", "Claude (Anthropic), ChatGPT, dan Google Flow (Veo 3.1).")
    add_p("Peranan & Pengesahan: ", "AI digunakan untuk pemetaan kurikulum NOSS, penyusunan skrip PdP, dan penjanaan segmen visual konsep. Semua kandungan teknikal, logik pengaturcaraan, dan pengujian fizikal prototaip disemak serta disahkan secara manual oleh penulis.")

    # --- 7. HASIL PENGUJIAN ---
    add_heading("7. HASIL PENGUJIAN / VALIDASI / BUKTI")
    add_p("Hasil pengujian rintis menunjukkan peningkatan kefahaman logik sistem ADAS yang ketara. Pelatih yang pada mulanya sukar membayangkan logik tindak balas AEB menerusi manual teknikal berjaya membuat deduksi logik diagnostik kerosakan pada mikropengawal secara automatik semasa simulasi kegagalan sensor.")

    # IMAGE 2: CHART
    img2_path = os.path.join('extracted_media', 'graf_hasil_pengujian.png')
    if os.path.exists(img2_path):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img2 = p_img2.add_run()
        run_img2.add_picture(img2_path, width=Inches(5.5))
        
        p_cap2 = doc.add_paragraph()
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap2 = p_cap2.add_run("Rajah 2. Carta Hasil Pengujian Rintis Perbandingan Kefahaman Pelatih Sebelum & Selepas Inovasi")
        r_cap2.font.italic = True
        r_cap2.font.bold = True
        r_cap2.font.size = Pt(9.5)
        r_cap2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # --- 8. IMPAK INOVASI ---
    add_heading("8. IMPAK INOVASI")
    
    table_impak = doc.add_table(rows=6, cols=2)
    table_impak.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_impak.autofit = False

    headers_i = ["Bidang Impak", "Penerangan"]
    hdr_cells_i = table_impak.rows[0].cells
    for i, title in enumerate(headers_i):
        hdr_cells_i[i].text = title
        set_cell_background(hdr_cells_i[i], "1E40AF")
        set_cell_margins(hdr_cells_i[i], 120, 120, 150, 150)
        p = hdr_cells_i[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10.5)

    data_impak = [
        ("Impak Akademik / Pembelajaran", "Merapatkan jurang teori dan praktikal logik kawalan ADAS sebelum latihan kenderaan sebenar."),
        ("Impak Operasi / Institusi", "Membantu pensyarah mencapai standard NOSS G452-011-4:2025 Level 4 secara kos efektif."),
        ("Impak Komuniti / Industri / Pasaran", "Menjadi kit latihan asas berpotensi tinggi untuk institusi kemahiran (Politeknik, ILP, ADTEC) dan sekolah vokasional."),
        ("Potensi Pengkomersialan", "Berpotensi dikomersialkan sebagai Kit Pendidikan STEM/TVET interaktif bersama pakej perkakasan dan perisian."),
        ("Impak Dasar / Ekosistem TVET", "Menyokong agenda Industri 4.0 dan transformasi TVET negara dengan pendedahan teknologi sensor AI kenderaan berautonomi.")
    ]

    for row_idx, data in enumerate(data_impak, start=1):
        row_cells = table_impak.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for c_idx in range(2):
            row_cells[c_idx].text = data[c_idx]
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], 100, 100, 150, 150)
            p = row_cells[c_idx].paragraphs[0]
            p.runs[0].font.size = Pt(10)
            if c_idx == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    for row in table_impak.rows:
        for i, w in enumerate(col_widths_n):
            row.cells[i].width = w

    # --- 9. KESIMPULAN ---
    add_heading("9. KESIMPULAN")
    add_p("SmartLine QR Rover berjaya membantu menyelesaikan masalah kekangan kos dan risiko keselamatan latihan diagnostik ADAS. Inovasi ini memberi manfaat tinggi dari aspek kefahaman logik kawalan dan berpotensi diperluaskan secara menyeluruh dalam ekosistem TVET Automotif.")

    # --- 10. RUJUKAN DAN BAHAN SOKONGAN ---
    add_heading("10. RUJUKAN DAN BAHAN SOKONGAN")
    add_p("Jabatan Pembangunan Kemahiran. (2025). National Occupational Skills Standard (NOSS) G452-011-4:2025 Automotive Electrical Diagnostic Level 4. Kementerian Sumber Manusia Malaysia.")
    
    add_p("Modul PdP Digital Google Site: ", "Imbas Kod QR di bawah untuk mengakses modul PdP interaktif, manual amali, dan kod sumber:")

    # IMAGE 3: QR CODE
    img3_path = os.path.join('extracted_media', 'qr_pdp_module.png')
    if os.path.exists(img3_path):
        p_img3 = doc.add_paragraph()
        p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img3 = p_img3.add_run()
        run_img3.add_picture(img3_path, width=Inches(2.2))
        
        p_cap3 = doc.add_paragraph()
        p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap3 = p_cap3.add_run("Kod QR Modul PdP Digital SmartLine QR Rover\n")
        r_cap3.font.bold = True
        r_cap3.font.size = Pt(10)
        
        r_url = p_cap3.add_run("Pautan: https://sites.google.com/jtm.gov.my/smartlineqrrover/home")
        r_url.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        r_url.font.underline = True
        r_url.font.size = Pt(10)

    out_docx = "Poster_SmartLine_QR_Rover_MIPACTVET2026.docx"
    doc.save(out_docx)
    print(f"Successfully generated poster document: {out_docx}")

if __name__ == "__main__":
    build_docx()
